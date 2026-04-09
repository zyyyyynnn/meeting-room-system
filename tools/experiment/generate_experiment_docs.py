from __future__ import annotations

import json
import math
import textwrap
from datetime import date
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENTATION, WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


REPO_ROOT = Path(__file__).resolve().parents[2]
DOCS_DIR = REPO_ROOT / "docs" / "experiment"
ASSETS_DIR = DOCS_DIR / "assets"
DRAWIO_SRC_DIR = ASSETS_DIR / "drawio-src"
TEMPLATE_PATH = Path(r"E:\私有云\课程设计格式.docx")
TODAY = f"{date.today().year}.{date.today().month}.{date.today().day}"
PROJECT_NAME = "会议室预约与资源协调系统"
PROJECT_ALIAS = "Meeting Room Reservation & Coordination System"
COVER_TITLE_MAP = {
    "需求规格说明书（SRS）": "会议室预约系统需求规格说明书（SRS）",
    "产品原型设计": "会议室预约系统产品原型设计",
    "用户故事地图": "会议室预约系统用户故事地图",
    "系统架构设计文档": "会议室预约系统架构设计",
    "数据库设计文档": "会议室预约系统数据库设计",
    "API 接口文档": "会议室预约系统 API 接口设计",
}


def pick_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf" if bold else r"C:\Windows\Fonts\simsun.ttc",
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size=size)
    return ImageFont.load_default()


def create_canvas(width: int, height: int) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    image = Image.new("RGB", (width, height), "white")
    return image, ImageDraw.Draw(image)


def wrap_cjk(text: str, width: int) -> list[str]:
    return textwrap.wrap(text, width=width, break_long_words=True, break_on_hyphens=False) or [text]


def draw_multiline_text(draw: ImageDraw.ImageDraw, xy: tuple[int, int], lines: Iterable[str], font, fill, line_gap: int = 8) -> None:
    x, y = xy
    for line in lines:
        draw.text((x, y), line, font=font, fill=fill)
        y += font.size + line_gap


def draw_box(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], title: str, body_lines: list[str], *, fill: str, outline: str) -> None:
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=3)
    draw.text((x1 + 18, y1 + 14), title, font=pick_font(26, bold=True), fill="black")
    draw.line((x1 + 16, y1 + 56, x2 - 16, y1 + 56), fill=outline, width=2)
    draw_multiline_text(draw, (x1 + 18, y1 + 74), body_lines, pick_font(18), fill="black", line_gap=6)


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], fill: str = "#8b0000", width: int = 4) -> None:
    sx, sy = start
    ex, ey = end
    draw.line((sx, sy, ex, ey), fill=fill, width=width)
    angle = math.atan2(ey - sy, ex - sx)
    arrow_len = 16
    spread = math.pi / 8
    left = (ex - arrow_len * math.cos(angle - spread), ey - arrow_len * math.sin(angle - spread))
    right = (ex - arrow_len * math.cos(angle + spread), ey - arrow_len * math.sin(angle + spread))
    draw.polygon([(ex, ey), left, right], fill=fill)


def create_architecture_diagram(path: Path) -> None:
    image, draw = create_canvas(1700, 920)
    draw.text((60, 40), "系统总体架构图", font=pick_font(40, bold=True), fill="#8b0000")
    draw.text((60, 95), "前后端分离 + JWT 鉴权 + MySQL 持久化 + Redis 运行态缓存", font=pick_font(22), fill="#444444")
    draw_box(draw, (60, 180, 390, 430), "浏览器端用户", ["PC 浏览器访问系统", "按角色浏览页面", "提交预约/审批操作"], fill="#f7efe8", outline="#b85c38")
    draw_box(draw, (460, 180, 860, 470), "Vue 3 前端", ["views + router + api + store", "Axios 调用接口", "Dashboard / Calendar / Rooms / Mine / Approvals"], fill="#edf6ff", outline="#3c78d8")
    draw_box(draw, (940, 120, 1430, 530), "Spring Boot 后端", ["Security + JWT 认证过滤链", "Auth / Reservation / Room / Stats / User 模块", "统一 ApiResponse 返回", "编排 Redis 位图、锁与通知"], fill="#edf7ed", outline="#3d8b3d")
    draw_box(draw, (1480, 120, 1640, 280), "MySQL 8", ["用户、会议室、预约", "权威业务数据"], fill="#fff4dd", outline="#c48a00")
    draw_box(draw, (1480, 340, 1640, 530), "Redis", ["位图缓存", "分布式锁", "通知/审计日志", "房间状态与维护信息"], fill="#fdebec", outline="#c94f5d")
    draw_box(draw, (460, 560, 860, 820), "前端关键能力", ["路由守卫：未登录跳转 /login", "登录成功统一进入 /dashboard", "后端不可达时展示状态面板", "通知抽屉统一查看消息"], fill="#faf4ff", outline="#7a4bb7")
    draw_box(draw, (940, 600, 1430, 820), "后端关键能力", ["预约时段校验（8:00-18:00 / 半小时粒度）", "审批与冲突检测", "房间状态/维护时段检查", "最近预约缓存与审计记录"], fill="#fff5f5", outline="#cc4b4b")
    draw_arrow(draw, (390, 305), (460, 305))
    draw_arrow(draw, (860, 305), (940, 305))
    draw_arrow(draw, (1430, 220), (1480, 220))
    draw_arrow(draw, (1430, 430), (1480, 430))
    draw_arrow(draw, (660, 470), (660, 560))
    draw_arrow(draw, (1185, 530), (1185, 600))
    image.save(path)


def create_auth_flow_diagram(path: Path) -> None:
    image, draw = create_canvas(1700, 860)
    draw.text((60, 40), "认证与鉴权流程图", font=pick_font(40, bold=True), fill="#8b0000")
    draw.text((60, 95), "登录页 -> AuthController -> JWT -> localStorage -> 鉴权过滤链", font=pick_font(22), fill="#444444")
    boxes = [
        (60, 220, 300, 370, "登录页 /login", ["输入用户名与密码", "可见服务状态面板"]),
        (360, 220, 620, 370, "POST /api/auth/login", ["AuthController 接收请求", "AuthService 校验凭据"]),
        (690, 220, 970, 400, "JWT 登录响应", ["返回 token / userId / username / role", "前端写入 mrs_auth"]),
        (1040, 180, 1360, 420, "Axios 请求拦截器", ["读取本地 token", "写入 Authorization 头"]),
        (1420, 180, 1640, 420, "SecurityFilterChain", ["JwtAuthFilter 解析令牌", "MethodSecurity 校验角色权限"]),
    ]
    for x1, y1, x2, y2, title, lines in boxes:
        draw_box(draw, (x1, y1, x2, y2), title, lines, fill="#f9fbff", outline="#4876d0")
    draw_box(draw, (560, 530, 1140, 760), "路由守卫与页面进入", ["未登录访问业务页 -> /login", "登录后访问认证页 -> /dashboard", "所有角色默认首页均为 /dashboard"], fill="#fff7e9", outline="#c48a00")
    draw_arrow(draw, (300, 295), (360, 295))
    draw_arrow(draw, (620, 295), (690, 295))
    draw_arrow(draw, (970, 295), (1040, 295))
    draw_arrow(draw, (1360, 295), (1420, 295))
    draw_arrow(draw, (830, 400), (830, 530))
    image.save(path)


def create_reservation_sequence_diagram(path: Path) -> None:
    image, draw = create_canvas(1800, 980)
    draw.text((60, 40), "预约创建时序图", font=pick_font(40, bold=True), fill="#8b0000")
    draw.text((60, 95), "用户发起预约时，前端、业务服务、Redis 与 MySQL 的协同过程", font=pick_font(22), fill="#444444")
    actors = ["用户", "CalendarView", "ReservationController", "ReservationService", "Redis", "MySQL", "Notification/Audit"]
    x_positions = [110, 340, 610, 900, 1180, 1420, 1660]
    for actor, x in zip(actors, x_positions):
        draw.text((x - 55, 170), actor, font=pick_font(22, bold=True), fill="#1f1f1f")
        draw.line((x, 220, x, 900), fill="#a0a0a0", width=2)
    steps = [
        (260, 110, 340, "选择房间与时段"),
        (330, 340, 610, "POST /api/reservations"),
        (400, 610, 900, "调用 create(req)"),
        (470, 900, 1180, "检查位图、房间状态、维护时段"),
        (540, 900, 1420, "校验冲突并写入 reservation"),
        (610, 1420, 900, "返回新增记录"),
        (680, 900, 1180, "rebuildDay / evictUserCache"),
        (750, 900, 1660, "推送通知并记录审计"),
        (820, 900, 610, "返回 ReservationResp"),
        (890, 610, 340, "返回 ApiResponse"),
    ]
    for y, sx, ex, label in steps:
        draw_arrow(draw, (sx, y), (ex, y))
        draw.rounded_rectangle((min(sx, ex) + 16, y - 24, max(sx, ex) - 16, y + 12), radius=12, fill="#fff7e9", outline="#c48a00", width=2)
        draw.text((min(sx, ex) + 28, y - 18), label, font=pick_font(18), fill="#333333")
    image.save(path)


def create_story_map_diagram(path: Path) -> None:
    image, draw = create_canvas(1900, 1100)
    draw.text((60, 40), "用户故事地图（主干活动 × 角色）", font=pick_font(38, bold=True), fill="#8b0000")
    draw.text((60, 92), "以角色为行、以主干活动为列组织当前系统的故事切片", font=pick_font(22), fill="#444444")
    activities = ["登录进入", "查看态势", "查找会议室", "发起预约", "管理个人预约", "审批预约", "会议室治理", "用户治理"]
    roles = ["USER", "ADMIN", "SUPER_ADMIN"]
    role_notes = {
        "USER": ["登录并进入看板", "查看与自己相关的资源态势", "查看日历并创建预约", "取消/删除自己的预约"],
        "ADMIN": ["聚合待审批与风险", "审批/驳回预约", "查看审批日志", "管理用户角色与启停状态"],
        "SUPER_ADMIN": ["继承 ADMIN 全能力", "维护房间状态和维护时段", "新增/编辑/删除会议室", "撤销审批结果重新审核"],
    }
    start_x, top_y, col_width, row_height = 220, 180, 200, 250
    for index, activity in enumerate(activities):
        x1 = start_x + index * col_width
        draw.rounded_rectangle((x1, top_y, x1 + col_width - 14, top_y + 70), radius=12, fill="#f4f8ff", outline="#3c78d8", width=3)
        draw_multiline_text(draw, (x1 + 16, top_y + 14), wrap_cjk(activity, 8), pick_font(22, bold=True), fill="#204a87", line_gap=4)
    for row, role in enumerate(roles):
        y1 = top_y + 100 + row * row_height
        draw.rounded_rectangle((60, y1, 190, y1 + row_height - 20), radius=14, fill="#fff2eb", outline="#b85c38", width=3)
        draw.text((80, y1 + 24), role, font=pick_font(24, bold=True), fill="#8b0000")
        draw_multiline_text(draw, (78, y1 + 70), wrap_cjk(" / ".join(role_notes[role]), 8), pick_font(16), fill="#333333", line_gap=4)
        for col, activity in enumerate(activities):
            x1 = start_x + col * col_width
            draw.rounded_rectangle((x1, y1, x1 + col_width - 14, y1 + row_height - 20), radius=10, fill="white", outline="#d0d7e2", width=2)
            mapping = {
                "USER": {"登录进入": "登录系统\n进入 /dashboard", "查看态势": "看板欢迎语\n个人提醒", "查找会议室": "筛选房间\n查看日历", "发起预约": "单次预约\n周期预约", "管理个人预约": "取消预约\n删除预约\n看个人日志"},
                "ADMIN": {"登录进入": "登录系统\n进入 /dashboard", "查看态势": "审批数量\n维护风险\n通知中心", "查找会议室": "查看日历\n判断资源", "发起预约": "可代入普通用户场景", "管理个人预约": "查看本人记录", "审批预约": "待审批列表\n批准/驳回\n审计日志", "用户治理": "查看用户\n改角色/启停"},
                "SUPER_ADMIN": {"登录进入": "登录系统\n进入 /dashboard", "查看态势": "高权限看板\n全局资源态势", "查找会议室": "查看日历\n可用资源", "发起预约": "可代入普通用户场景", "管理个人预约": "查看本人记录", "审批预约": "批准/驳回\n撤销审批", "会议室治理": "新增/编辑/删除\n状态维护\n维护时段", "用户治理": "继承管理员治理能力"},
            }
            content = mapping[role].get(activity, "—")
            draw_multiline_text(draw, (x1 + 12, y1 + 18), wrap_cjk(content, 10), pick_font(18), fill="#333333", line_gap=4)
    image.save(path)


def create_database_er_diagram(path: Path) -> None:
    image, draw = create_canvas(1700, 950)
    draw.text((60, 40), "MySQL 核心表 ER 图", font=pick_font(40, bold=True), fill="#8b0000")
    draw.text((60, 95), "当前版本以 sys_user、meeting_room、reservation 为核心业务表", font=pick_font(22), fill="#444444")
    draw_box(draw, (80, 200, 520, 670), "sys_user", ["PK id : bigint", "username : varchar(64) UNIQUE", "password_hash : varchar(255)", "role : varchar(32)", "enabled : tinyint", "created_at / updated_at : datetime"], fill="#f9fbff", outline="#3c78d8")
    draw_box(draw, (620, 160, 1110, 780), "reservation", ["PK id : bigint", "FK user_id -> sys_user.id", "FK room_id -> meeting_room.id", "start_time / end_time : datetime", "status : varchar(32)", "reason / admin_comment : varchar(255)", "approved_by / approved_at", "INDEX idx_room_time(room_id,start_time,end_time)", "INDEX idx_user_time(user_id,start_time)"], fill="#fff7e9", outline="#c48a00")
    draw_box(draw, (1210, 200, 1630, 670), "meeting_room", ["PK id : bigint", "name : varchar(128) UNIQUE", "capacity : int", "equipment_json : json", "require_approval : tinyint", "created_at / updated_at : datetime"], fill="#edf7ed", outline="#3d8b3d")
    draw_arrow(draw, (520, 420), (620, 420))
    draw_arrow(draw, (1210, 420), (1110, 420))
    image.save(path)


def set_cell_text(cell, text: str, *, bold: bool = False, size: float = 10.5, color: RGBColor | None = None, align: int = WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if color:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_page(doc: Document, *, landscape: bool = False) -> None:
    section = doc.sections[0]
    section.page_width = Cm(29.7 if landscape else 21)
    section.page_height = Cm(21 if landscape else 29.7)
    section.orientation = WD_ORIENTATION.LANDSCAPE if landscape else WD_ORIENTATION.PORTRAIT
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.2)


def set_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(6)
    for style_name, size, color in (("Heading 1", 16, "8B0000"), ("Heading 2", 13, "1F4E79"), ("Heading 3", 11.5, "333333")):
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def create_document(title: str, landscape: bool = False) -> Document:
    doc = Document()
    set_page(doc, landscape=landscape)
    set_styles(doc)
    doc.core_properties.title = title
    doc.core_properties.subject = PROJECT_NAME
    doc.core_properties.author = "Codex"
    doc.core_properties.comments = "依据当前代码与运行结果生成"
    add_cover(doc, title)
    return doc


def add_cover(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(70)
    run = p.add_run(PROJECT_NAME)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("8B0000")
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(title)
    run2.font.name = "Microsoft YaHei"
    run2._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run2.font.size = Pt(20)
    run2.font.bold = True
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.space_before = Pt(12)
    p3.add_run(PROJECT_ALIAS)
    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.style = "Table Grid"
    rows = [("交付日期", TODAY), ("文档说明", "基于当前仓库代码、接口与运行页面生成"), ("交付格式", "仅保留 .docx 正式文档"), ("交付目录", str(DOCS_DIR))]
    for row_idx, (label, value) in enumerate(rows):
        set_cell_text(meta.cell(row_idx, 0), label, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(meta.cell(row_idx, 1), value)
        shade_cell(meta.cell(row_idx, 0), "F3F6FA")
    doc.add_page_break()


def add_paragraph(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_bullet_list(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[Iterable[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(table.rows[0].cells[idx], "DCE6F1")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], str(value))


def add_image(doc: Document, path: Path, caption: str, width_cm: float = 16.0) -> None:
    if not path.exists():
        doc.add_paragraph(f"图资源缺失：{path.name}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.color.rgb = RGBColor.from_string("666666")


def add_code_block(doc: Document, data: dict | list | None) -> None:
    if data is None:
        doc.add_paragraph("无")
        return
    for line in json.dumps(data, ensure_ascii=False, indent=2).splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9.5)


def create_diagram_assets() -> None:
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    DRAWIO_SRC_DIR.mkdir(parents=True, exist_ok=True)
    targets = [
        (ASSETS_DIR / "diagram-architecture-overview.png", create_architecture_diagram),
        (ASSETS_DIR / "diagram-auth-flow.png", create_auth_flow_diagram),
        (ASSETS_DIR / "diagram-reservation-sequence.png", create_reservation_sequence_diagram),
        (ASSETS_DIR / "diagram-story-map.png", create_story_map_diagram),
        (ASSETS_DIR / "diagram-database-er.png", create_database_er_diagram),
    ]
    for target, builder in targets:
        if not target.exists():
            builder(target)


STACK_ROWS = [
    ("前端", "Vue 3、TypeScript、Vite、Element Plus、FullCalendar", "负责页面呈现、路由守卫、接口调用与交互反馈"),
    ("后端", "Spring Boot 3、Spring Security、JWT、MyBatis-Plus", "负责认证鉴权、业务编排、权限控制和数据访问"),
    ("关系数据库", "MySQL 8", "保存用户、会议室、预约等核心业务数据"),
    ("缓存/运行态", "Redis", "保存预约位图、分布式锁、通知、审计日志、房间状态与维护信息"),
]

ROLE_ROWS = [
    ("USER", "普通用户", "登录、查看个人看板、按日历预约会议室、周期预约、查看/取消/删除自己的预约、接收通知"),
    ("ADMIN", "管理员", "具备 USER 全部能力，并可处理预约审批、查看审批审计、进入用户管理页面"),
    ("SUPER_ADMIN", "超级管理员", "具备 ADMIN 全部能力，并可撤销审批结果、创建/编辑/删除会议室、维护房间状态与维护时段"),
]

FEATURE_ROWS = [
    ("FR-01", "认证与会话", "登录、注册、JWT 会话保持、路由守卫、失效后跳转登录", "游客、所有已登录用户"),
    ("FR-02", "运营看板", "基于角色展示欢迎语、任务摘要、资源态势、风险分布与快捷入口", "USER / ADMIN / SUPER_ADMIN"),
    ("FR-03", "会议预约日历", "按会议室查看日历、查询日占用、识别冲突并给出替代建议", "所有已登录用户"),
    ("FR-04", "预约创建", "创建单次预约、创建每周重复预约、按审批规则自动置为待审批或已批准", "所有已登录用户"),
    ("FR-05", "我的预约", "查看近 30 天预约、取消预约、删除预约、查看个人审计记录", "所有已登录用户"),
    ("FR-06", "预约审批", "查看待审批队列、批准、驳回、查看最近审批结果、查看全局审计日志", "ADMIN / SUPER_ADMIN"),
    ("FR-07", "会议室治理", "浏览会议室、维护容量/设备、变更房间状态、登记维护时段、删除房间", "SUPER_ADMIN"),
    ("FR-08", "用户管理", "查看账号列表、修改角色、启停账号", "ADMIN / SUPER_ADMIN"),
    ("FR-09", "通知中心", "汇总预约、审批、会议提醒等通知，从右侧抽屉集中查看", "所有已登录用户"),
]

BUSINESS_RULE_ROWS = [
    ("BR-01", "预约时间粒度", "当前版本按半小时为最小时间单位，分钟仅允许 00 或 30。"),
    ("BR-02", "可预约时间窗", "只支持 08:00-18:00 的工作时段，且必须在同一天内完成预约。"),
    ("BR-03", "单次时长限制", "单次预约时长不超过 120 分钟。"),
    ("BR-04", "提前预约范围", "只允许预约未来 7 天内的会议。"),
    ("BR-05", "单人日上限", "单个用户在同一天最多保留 2 条处于 PENDING/APPROVED 的预约。"),
    ("BR-06", "周期预约", "每周周期预约仅允许 1-12 周。"),
    ("BR-07", "阻塞状态", "PENDING 和 APPROVED 都会阻塞同会议室同时间段的后续预约与审批。"),
    ("BR-08", "会议室可用性", "只有状态为 AVAILABLE 且不处于维护时段的会议室才允许创建预约。"),
    ("BR-09", "审批规则", "房间 requireApproval=true 时默认走审批；A-101 作为特例强制审批。"),
    ("BR-10", "状态流转", "预约状态支持 PENDING、APPROVED、REJECTED、CANCELLED。"),
]

NON_FUNCTIONAL_ROWS = [
    ("NFR-01", "安全性", "使用 BCrypt 存储密码哈希；接口通过 JWT 进行无状态鉴权；敏感操作使用角色校验。"),
    ("NFR-02", "可用性", "前端具备后端健康状态面板；接口统一返回 ApiResponse；后端使用 Redis 锁避免并发抢占。"),
    ("NFR-03", "性能", "日占用信息通过 Redis 位图缓存；我的预约通过 Redis 缓存最近 30 天数据；数据库使用复合索引支撑查询。"),
    ("NFR-04", "可维护性", "前端按 views/api/store/router 分层；后端按 auth/room/reservation/stats/user 模块化组织。"),
    ("NFR-05", "兼容性", "开发环境面向 Windows，前端基于现代浏览器，后端运行在 Java 21 / Maven 3.8+。"),
]

PAGE_ROWS = [
    ("登录页", "/login", "游客", "输入用户名和密码，完成系统入口认证，并展示后端联通状态"),
    ("运营看板", "/dashboard", "所有角色", "以角色为中心展示待办、资源态势、趋势和快捷入口"),
    ("会议预约日历", "/calendar", "所有角色", "查看会议室日历、日占用和冲突建议"),
    ("会议室管理", "/rooms", "SUPER_ADMIN（治理操作）", "维护会议室容量、设备、状态和维护时段"),
    ("我的预约", "/mine", "所有角色", "查看个人预约、取消/删除预约、查看个人审计记录"),
    ("预约审批", "/admin/approvals", "ADMIN / SUPER_ADMIN", "处理待审批预约，超级管理员可撤销审批结果"),
]

PAGE_DETAILS = [
    ("登录页", ASSETS_DIR / "prototype-01-login.png", "/login", "游客", "完成系统身份认证，并在异常情况下给出后端服务状态提示。", "品牌视觉区、用户名/密码输入区、登录按钮、注册链接、服务状态面板。", "输入合法凭据后调用 /api/auth/login；成功后根据角色统一进入 /dashboard；当后端不可达时在页内展示恢复提示。"),
    ("运营看板", ASSETS_DIR / "prototype-02-dashboard.png", "/dashboard", "所有角色", "聚合今日待办、资源态势和风险信号，作为进入其他业务页面的中枢。", "欢迎区、任务摘要卡片、资源快照、热力趋势、风险分布、快捷入口。", "点击任务卡可带上下文跳转；支持手动刷新；管理员与普通用户看到不同的指标聚合。"),
    ("会议预约日历", ASSETS_DIR / "prototype-03-calendar.png", "/calendar", "所有角色", "以日历视图完成会议室时段判断，并辅助用户规避冲突。", "会议室筛选、FullCalendar 日历、时段占用条、冲突提示、替代方案面板。", "选择房间和时间范围后查询 /api/reservations/calendar；如冲突则联动 /suggestions 提供同房间顺延或其他房间方案。"),
    ("会议室管理", ASSETS_DIR / "prototype-04-rooms.png", "/rooms", "SUPER_ADMIN（治理操作）", "统一治理会议室主数据、状态和维护计划。", "统计卡、筛选区、会议室列表、编辑对话框、状态维护、维护时段录入。", "支持新增/编辑/删除会议室；支持切换 AVAILABLE / MAINTENANCE / DISABLED；维护时段写入 Redis 列表。"),
    ("我的预约", ASSETS_DIR / "prototype-05-mine.png", "/mine", "所有角色", "集中管理个人近 30 天预约，并保留用户侧审计轨迹。", "预约统计、预约列表、取消/删除操作、个人审计记录区。", "读取 /api/reservations/mine/recent；针对自己的预约执行取消或删除；同步查看 Redis 缓存的用户审计日志。"),
    ("预约审批", ASSETS_DIR / "prototype-06-approvals.png", "/admin/approvals", "ADMIN / SUPER_ADMIN", "管理员处理待审批预约，超级管理员可撤销历史审批结果并重审。", "待审批列表、最近已审批记录、全局审计日志、批准/驳回/撤销操作。", "待审批列表来自 /api/admin/reservations/pending；批准和驳回会写回审批意见并触发通知；SUPER_ADMIN 可执行撤销审批。"),
]

STORY_ROWS = [
    ("US-01", "USER", "登录进入", "作为普通用户，我希望登录后直接进入运营看板，以便快速开始当天预约。", "P0"),
    ("US-02", "USER", "查看态势", "作为普通用户，我希望在看板上看到与我相关的预约提醒和资源概览，以便判断是否需要发起会议。", "P0"),
    ("US-03", "USER", "查找会议室", "作为普通用户，我希望按会议室和时间查看日历，以便找到可用时段。", "P0"),
    ("US-04", "USER", "发起预约", "作为普通用户，我希望创建单次预约，以便锁定具体会议时间。", "P0"),
    ("US-05", "USER", "发起预约", "作为普通用户，我希望创建每周重复预约，以便处理例会场景。", "P1"),
    ("US-06", "USER", "管理个人预约", "作为普通用户，我希望取消或删除自己的预约，以便及时释放资源。", "P0"),
    ("US-07", "USER", "管理个人预约", "作为普通用户，我希望查看预约相关通知和个人审计记录，以便追踪关键动作。", "P1"),
    ("US-08", "ADMIN", "查看态势", "作为管理员，我希望在看板中看到待审批数量和维护风险，以便优先处理异常。", "P0"),
    ("US-09", "ADMIN", "审批预约", "作为管理员，我希望批准符合条件的预约，以便让用户尽快完成协同安排。", "P0"),
    ("US-10", "ADMIN", "审批预约", "作为管理员，我希望驳回不符合条件的预约并填写备注，以便说明原因。", "P0"),
    ("US-11", "ADMIN", "审批预约", "作为管理员，我希望查看最近已审批记录和全局审计日志，以便追溯审批动作。", "P1"),
    ("US-12", "ADMIN", "用户治理", "作为管理员，我希望查看用户列表并调整角色/启停状态，以便维护账号秩序。", "P1"),
    ("US-13", "SUPER_ADMIN", "会议室治理", "作为超级管理员，我希望新增或编辑会议室，以便让系统资源配置保持最新。", "P0"),
    ("US-14", "SUPER_ADMIN", "会议室治理", "作为超级管理员，我希望维护房间状态和维护时段，以便避免错误预约。", "P0"),
    ("US-15", "SUPER_ADMIN", "会议室治理", "作为超级管理员，我希望删除没有阻塞预约的会议室，以便清理无效资源。", "P2"),
    ("US-16", "SUPER_ADMIN", "审批预约", "作为超级管理员，我希望撤销已审批结果并重新审核，以便纠正误审批。", "P1"),
    ("US-17", "SUPER_ADMIN", "用户治理", "作为超级管理员，我希望继承管理员的用户管理能力，以便完成更高权限治理。", "P1"),
]

MYSQL_TABLES = {
    "sys_user": {
        "purpose": "保存系统用户、密码哈希、角色和启停状态。",
        "fields": [
            ("id", "bigint", "PK, AUTO_INCREMENT", "用户主键"),
            ("username", "varchar(64)", "NOT NULL, UNIQUE", "登录用户名"),
            ("password_hash", "varchar(255)", "NOT NULL", "BCrypt 密码哈希"),
            ("role", "varchar(32)", "NOT NULL", "角色，取值 USER / ADMIN / SUPER_ADMIN"),
            ("enabled", "tinyint", "NOT NULL, DEFAULT 1", "账号是否启用"),
            ("created_at", "datetime", "NOT NULL, DEFAULT CURRENT_TIMESTAMP", "创建时间"),
            ("updated_at", "datetime", "NOT NULL, DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP", "更新时间"),
        ],
        "indexes": [("PRIMARY", "id", "主键索引"), ("username", "username", "用户名唯一索引")],
    },
    "meeting_room": {
        "purpose": "保存会议室主数据、容量、设备 JSON 和审批配置。",
        "fields": [
            ("id", "bigint", "PK, AUTO_INCREMENT", "会议室主键"),
            ("name", "varchar(128)", "NOT NULL, UNIQUE", "会议室名称"),
            ("capacity", "int", "NOT NULL", "容纳人数"),
            ("equipment_json", "json", "NULL", "设备列表 JSON"),
            ("require_approval", "tinyint", "NOT NULL, DEFAULT 0", "是否需要审批"),
            ("created_at", "datetime", "NOT NULL, DEFAULT CURRENT_TIMESTAMP", "创建时间"),
            ("updated_at", "datetime", "NOT NULL, DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP", "更新时间"),
        ],
        "indexes": [("PRIMARY", "id", "主键索引"), ("name", "name", "会议室名称唯一索引")],
    },
    "reservation": {
        "purpose": "保存预约明细、审批信息以及用户与会议室的关联关系。",
        "fields": [
            ("id", "bigint", "PK, AUTO_INCREMENT", "预约主键"),
            ("user_id", "bigint", "NOT NULL, FK -> sys_user.id", "预约人"),
            ("room_id", "bigint", "NOT NULL, FK -> meeting_room.id", "会议室"),
            ("start_time", "datetime", "NOT NULL", "开始时间"),
            ("end_time", "datetime", "NOT NULL", "结束时间"),
            ("status", "varchar(32)", "NOT NULL", "状态，PENDING / APPROVED / REJECTED / CANCELLED"),
            ("reason", "varchar(255)", "NULL", "预约原因"),
            ("admin_comment", "varchar(255)", "NULL", "审批备注"),
            ("approved_by", "bigint", "NULL", "审批人 ID（当前未建立外键）"),
            ("approved_at", "datetime", "NULL", "审批时间"),
            ("created_at", "datetime", "NOT NULL, DEFAULT CURRENT_TIMESTAMP", "创建时间"),
            ("updated_at", "datetime", "NOT NULL, DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP", "更新时间"),
        ],
        "indexes": [
            ("PRIMARY", "id", "主键索引"),
            ("idx_room_time", "room_id, start_time, end_time", "会议室时间段查询索引"),
            ("idx_user_time", "user_id, start_time", "用户近 30 天预约查询索引"),
        ],
    },
}

REDIS_ROWS = [
    ("mrs:room:{roomId}:day:{yyyyMMdd}:hour-bitmap", "String bitmap", "缓存房间某日半小时粒度的占用位图"),
    ("mrs:room:{roomId}:day:{yyyyMMdd}:slices", "String(JSON)", "缓存房间某日的时段切片明细"),
    ("mrs:lock:room:{roomId}:day:{yyyyMMdd}", "String", "同房间同日期预约的分布式锁"),
    ("mrs:user:{userId}:reservations:last30d", "String(JSON)", "缓存用户最近 30 天预约列表"),
    ("mrs:user:{userId}:notifications", "List", "用户通知列表，最多保留 50 条"),
    ("mrs:user:{userId}:audit", "List", "用户审计日志"),
    ("mrs:admin:audit", "List", "管理员侧全局审计日志"),
    ("mrs:room:{roomId}:status", "String", "会议室运行状态 AVAILABLE / MAINTENANCE / DISABLED"),
    ("mrs:room:{roomId}:maintenances", "List(JSON)", "会议室维护时段列表"),
]

API_SUMMARY_ROWS = [
    ("认证", "POST", "/api/auth/register", "公开", "用户注册"),
    ("认证", "POST", "/api/auth/login", "公开", "用户登录并返回 JWT"),
    ("会议室", "GET", "/api/rooms", "已登录", "获取会议室列表"),
    ("会议室", "POST", "/api/rooms", "SUPER_ADMIN", "新增会议室"),
    ("会议室", "PUT", "/api/rooms/{id}", "SUPER_ADMIN", "更新会议室"),
    ("会议室", "DELETE", "/api/rooms/{id}", "SUPER_ADMIN", "删除会议室"),
    ("会议室", "POST", "/api/rooms/{id}/status", "SUPER_ADMIN", "修改会议室状态"),
    ("会议室", "POST", "/api/rooms/{id}/maintenances", "SUPER_ADMIN", "新增维护时段"),
    ("预约", "POST", "/api/reservations", "已登录", "创建单次预约"),
    ("预约", "POST", "/api/reservations/batch/weekly", "已登录", "创建每周重复预约"),
    ("预约", "POST", "/api/reservations/{id}/cancel", "本人", "取消自己的预约"),
    ("预约", "DELETE", "/api/reservations/{id}", "本人", "删除自己的预约"),
    ("预约", "GET", "/api/reservations/mine/recent", "已登录", "查询近 30 天个人预约"),
    ("预约", "GET", "/api/reservations/calendar", "已登录", "查询某会议室日历区间预约"),
    ("预约", "GET", "/api/reservations/occupancy/day", "已登录", "查询房间某日占用情况"),
    ("预约", "GET", "/api/reservations/suggestions", "已登录", "查询冲突替代方案"),
    ("审批", "GET", "/api/admin/reservations/pending", "ADMIN / SUPER_ADMIN", "待审批预约列表"),
    ("审批", "GET", "/api/admin/reservations/recent", "ADMIN / SUPER_ADMIN", "最近审批记录"),
    ("审批", "POST", "/api/admin/reservations/{id}/approve", "ADMIN / SUPER_ADMIN", "批准预约"),
    ("审批", "POST", "/api/admin/reservations/{id}/reject", "ADMIN / SUPER_ADMIN", "驳回预约"),
    ("审批", "POST", "/api/admin/reservations/{id}/revoke", "SUPER_ADMIN", "撤销审批结果"),
    ("统计", "GET", "/api/stats/overview", "已登录", "概览统计"),
    ("统计", "GET", "/api/stats/dashboard", "已登录", "运营看板数据"),
    ("用户管理", "GET", "/api/admin/users", "ADMIN / SUPER_ADMIN", "用户列表"),
    ("用户管理", "POST", "/api/admin/users/{id}/role", "ADMIN / SUPER_ADMIN", "修改用户角色"),
    ("用户管理", "POST", "/api/admin/users/{id}/enabled", "ADMIN / SUPER_ADMIN", "修改用户启停状态"),
]

API_DETAILS = [
    ("POST /api/auth/login", "校验用户名和密码，返回 JWT 及当前用户角色信息。", "公开接口", [("username", "string", "是", "用户名"), ("password", "string", "是", "密码")], {"username": "root", "password": "Root@123456"}, {"code": 0, "message": "OK", "data": {"token": "<JWT>", "userId": 3, "username": "root", "role": "SUPER_ADMIN"}}),
    ("GET /api/rooms", "返回会议室列表、状态、设备和维护时段。", "任意已登录用户", [], None, {"code": 0, "message": "OK", "data": [{"id": 1, "name": "A-101", "capacity": 8, "equipment": ["白板", "投影仪", "音响"], "requireApproval": False, "status": "AVAILABLE", "maintenanceSlots": []}]}),
    ("POST /api/reservations", "创建单次预约，并按房间审批策略决定初始状态。", "任意已登录用户", [("roomId", "number", "是", "会议室 ID"), ("startTime", "LocalDateTime", "是", "开始时间，半小时粒度"), ("endTime", "LocalDateTime", "是", "结束时间，半小时粒度"), ("reason", "string", "否", "预约原因")], {"roomId": 5, "startTime": "2026-04-09T09:00:00", "endTime": "2026-04-09T10:00:00", "reason": "项目周会"}, {"code": 0, "message": "OK", "data": {"id": 105, "userId": 3, "username": "root", "roomId": 5, "roomName": "北辰厅", "startTime": "2026-04-09T09:00:00", "endTime": "2026-04-09T10:00:00", "status": "PENDING", "reason": "项目周会", "adminComment": None, "approvedBy": None, "approvedAt": None}}),
    ("GET /api/reservations/suggestions", "在目标时段冲突时，提供同房间顺延/提前或其他房间候选方案。", "任意已登录用户", [("roomId", "number", "是", "会议室 ID"), ("start", "LocalDateTime", "是", "预约开始时间"), ("end", "LocalDateTime", "是", "预约结束时间")], {"roomId": 1, "start": "2026-04-09T09:00:00", "end": "2026-04-09T10:00:00"}, {"code": 0, "message": "OK", "data": {"conflictMessage": "当前时段与已有预约冲突，已为你推荐替代方案", "alternatives": [{"roomId": 1, "roomName": "A-101", "startTime": "2026-04-09T10:00:00", "endTime": "2026-04-09T11:00:00", "tip": "同会议室顺延"}]}}),
    ("GET /api/admin/reservations/pending", "查询待审批预约队列。", "ADMIN / SUPER_ADMIN", [], None, {"code": 0, "message": "OK", "data": [{"id": 27, "userId": 1, "username": "admin", "roomId": 1, "roomName": "A-101", "startTime": "2026-03-28T17:30:00", "endTime": "2026-03-28T18:00:00", "status": "PENDING", "reason": "在线预约（A-101）", "adminComment": None, "approvedBy": None, "approvedAt": None}]}),
    ("POST /api/admin/reservations/{id}/approve", "批准待审批预约并记录审批备注。", "ADMIN / SUPER_ADMIN", [("adminComment", "string", "否", "审批备注")], {"adminComment": "会议内容明确，批准执行"}, {"code": 0, "message": "OK", "data": None}),
    ("GET /api/stats/dashboard", "返回看板聚合数据，包括任务摘要、资源态势、趋势与快捷入口。", "任意已登录用户", [], None, {"code": 0, "message": "OK", "data": {"adminView": True, "welcome": {"roleLabel": "超级管理员", "message": "Welcome，root！"}, "taskSummary": {"title": "今日待办与异常", "subtitle": "先处理待审批、维护异常和近时段会议，再进入下方图层判断整体资源压力。", "items": [{"key": "pending-approvals", "label": "待审批预约", "value": 8, "detail": "需要优先人工处理", "tone": "warning", "to": "/admin/approvals", "query": {"focus": "pending"}}]}}}),
    ("GET /api/admin/users", "获取用户列表和角色/启停状态。", "ADMIN / SUPER_ADMIN", [], None, {"code": 0, "message": "OK", "data": [{"id": 1, "username": "admin", "role": "ADMIN", "enabled": True}]}),
]


def build_srs() -> Path:
    doc = create_document("需求规格说明书（SRS）")
    doc.add_heading("1. 项目概述", level=1)
    add_paragraph(doc, "本系统面向企业或团队会议室资源协调场景，采用前后端分离架构，实现登录注册、运营看板、会议预约日历、会议室治理、我的预约、预约审批、用户管理和通知中心等能力。")
    add_paragraph(doc, "当前版本的核心目标是：统一管理会议室主数据，降低预约冲突，提高审批可视性，并通过看板帮助不同角色快速进入当天的工作上下文。")
    doc.add_heading("2. 技术与运行环境", level=1)
    add_table(doc, ["层次", "技术选型", "说明"], STACK_ROWS)
    add_table(doc, ["环境项", "要求"], [("操作系统", "Windows（当前开发与运行环境）"), ("Java", "Java 21"), ("Maven", "3.8+"), ("Node.js", "18+"), ("数据库", "MySQL 8"), ("缓存", "Redis")])
    doc.add_heading("3. 角色与权限", level=1)
    add_table(doc, ["角色编码", "角色名称", "当前版本权限范围"], ROLE_ROWS)
    doc.add_heading("4. 功能需求", level=1)
    add_table(doc, ["编号", "模块", "需求描述", "适用角色"], FEATURE_ROWS)
    doc.add_heading("5. 业务规则", level=1)
    add_table(doc, ["编号", "规则名称", "规则说明"], BUSINESS_RULE_ROWS)
    doc.add_heading("6. 非功能需求", level=1)
    add_table(doc, ["编号", "维度", "当前要求"], NON_FUNCTIONAL_ROWS)
    doc.add_heading("7. 约束与边界", level=1)
    add_bullet_list(doc, ["当前版本仅支持单日内预约，不支持跨日会议。", "预约最小粒度固定为 30 分钟，适合会议室标准排班场景。", "会议室治理型接口由后端方法级权限控制，核心治理操作仅 SUPER_ADMIN 可调用。", "通知、审计日志和部分运行态信息保存在 Redis 中，不属于 MySQL 主表结构的一部分。", "本说明书以当前仓库代码、运行页面和数据库结构为事实依据，不扩展未实现功能。"])
    output = DOCS_DIR / "01-需求规格说明书.docx"
    return safe_save(doc, output)


def build_prototype_doc() -> Path:
    doc = create_document("产品原型设计")
    doc.add_heading("1. 原型说明", level=1)
    add_paragraph(doc, "本原型文档直接基于当前系统真实页面截图整理，不额外绘制 Axure/Figma 原型。截图来源于本地运行的 Vue 前端页面，能够对应当前仓库中的真实路由、组件和接口数据。")
    add_table(doc, ["页面", "路由", "主要角色", "页面价值"], PAGE_ROWS)
    doc.add_heading("2. 页面导航关系", level=1)
    add_bullet_list(doc, ["统一登录入口：/login；登录成功后按角色统一进入 /dashboard。", "顶部导航覆盖运营看板、会议预约、我的预约、会议室、预约审批和用户管理。", "通知中心通过右侧抽屉呈现，不单独占用路由页面。"])
    doc.add_heading("3. 核心页面原型", level=1)
    for index, (title, image, route, role, goal, modules, interactions) in enumerate(PAGE_DETAILS, start=1):
        doc.add_heading(f"3.{index} {title}", level=2)
        add_image(doc, image, f"图 {index} - {title} 页面截图", width_cm=16.4)
        add_table(doc, ["维度", "说明"], [("页面路由", route), ("适用角色", role), ("页面目标", goal), ("核心模块", modules), ("关键交互", interactions)])
    output = DOCS_DIR / "02-产品原型设计.docx"
    return safe_save(doc, output)


def build_story_map_doc() -> Path:
    doc = Document(TEMPLATE_PATH) if TEMPLATE_PATH.exists() else Document()
    set_page(doc, landscape=False, section_index=0)
    set_styles(doc)
    doc.core_properties.title = "用户故事地图"
    doc.core_properties.subject = PROJECT_NAME
    doc.core_properties.author = "Codex"
    doc.core_properties.comments = "课程设计文档"
    add_cover(doc, "用户故事地图")
    if TEMPLATE_PATH.exists():
        while len(doc.paragraphs) > 12:
            delete_paragraph(doc.paragraphs[-1])
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    set_page(doc, landscape=False, section_index=len(doc.sections) - 1)
    doc.add_heading("1. 故事地图概览", level=1)
    add_paragraph(doc, "本故事地图按 backbone、activities、stories 和 priority slices 组织当前版本能力：顶部展示主干活动，第二层展示关键任务，横向切片展示实验核心、当前增强与治理扩展；每张故事卡均用 USER、ADMIN、SUPER_ADMIN 标签标识适用角色。")
    add_image(doc, ASSETS_DIR / "diagram-story-map.png", "图 1 - 会议室预约系统分层用户故事地图", width_cm=14.8)
    doc.add_heading("2. 用户故事清单", level=1)
    add_table(doc, ["编号", "角色", "活动", "用户故事", "优先级"], STORY_ROWS)
    doc.add_heading("3. 优先级说明", level=1)
    add_bullet_list(doc, ["P0：直接支撑预约主流程和审批闭环，属于本实验必须覆盖内容。", "P1：提升治理效率或可追溯性，属于当前版本已实现的重要增强能力。", "P2：治理型扩展能力，当前版本已具备基础实现但使用频次相对较低。"])
    output = DOCS_DIR / "03-用户故事地图.docx"
    return safe_save(doc, output)


def build_architecture_doc() -> Path:
    doc = create_document("系统架构设计文档")
    doc.add_heading("1. 架构目标", level=1)
    add_bullet_list(doc, ["通过前后端分离保持界面迭代与业务服务解耦。", "通过 JWT 与 Spring Security 实现无状态鉴权。", "通过 MySQL 持久化业务主数据，通过 Redis 加速高频读写与运行态计算。", "通过模块化服务拆分认证、预约、审批、会议室治理、统计与用户治理能力。"])
    doc.add_heading("2. 总体架构", level=1)
    add_image(doc, ASSETS_DIR / "diagram-architecture-overview.png", "图 1 - 系统总体架构图")
    add_table(doc, ["层次", "说明"], [("表示层", "Vue 3 单页应用，负责页面渲染、路由守卫、接口调用和错误反馈。"), ("应用层", "Spring Boot 控制器与服务层，负责请求接入、业务编排、权限判断。"), ("数据层", "MyBatis-Plus 访问 MySQL；Redis 保存缓存、锁和运行态信息。"), ("安全层", "SecurityFilterChain + JwtAuthFilter + 方法级权限注解。")])
    doc.add_heading("3. 认证与鉴权设计", level=1)
    add_image(doc, ASSETS_DIR / "diagram-auth-flow.png", "图 2 - 认证与鉴权流程图")
    add_bullet_list(doc, ["登录成功后，前端把 token、userId、username、role 序列化到 localStorage 的 mrs_auth 键中。", "Axios 请求拦截器从本地会话中取 token，统一追加 Authorization 请求头。", "后端通过 JwtAuthFilter 解析 token，再交给方法级权限注解做角色控制。", "路由层在未登录时强制回到 /login；登录后再访问认证页会重定向到 /dashboard。"])
    doc.add_heading("4. 预约处理设计", level=1)
    add_image(doc, ASSETS_DIR / "diagram-reservation-sequence.png", "图 3 - 预约创建时序图")
    add_bullet_list(doc, ["ReservationService 会在真正写库前校验时间粒度、未来天数、时长上限、会议室状态和维护时段。", "同房间同日期的预约创建使用 Redis 分布式锁保护，避免并发抢占导致重复写入。", "预约创建和审批通过后都会重建当日位图缓存，并淘汰用户最近预约缓存。", "通知和审计由独立服务落到 Redis，支撑通知抽屉与日志追踪。"])
    doc.add_heading("5. 模块职责划分", level=1)
    add_table(doc, ["模块", "主要职责"], [("auth", "注册、登录、返回 JWT 与角色信息"), ("security", "JwtAuthFilter、JwtService、SecurityConfig、SecurityUtil"), ("reservation", "预约创建、取消、删除、冲突建议、审批、通知、审计"), ("room", "会议室主数据、状态、维护时段治理"), ("stats", "概览统计与运营看板聚合"), ("user", "用户列表、角色修改、账号启停"), ("bootstrap", "管理员/演示数据初始化")])
    output = DOCS_DIR / "04-系统架构设计文档.docx"
    return safe_save(doc, output)


def build_database_doc() -> Path:
    doc = create_document("数据库设计文档")
    doc.add_heading("1. 存储设计概述", level=1)
    add_paragraph(doc, "当前系统采用“ MySQL 负责权威业务数据 + Redis 负责运行态与缓存 ”的组合。MySQL 仅包含 3 张核心业务表，Redis 用于预约位图、锁、通知、审计和房间状态等短路径数据。")
    doc.add_heading("2. MySQL 核心表设计", level=1)
    add_image(doc, ASSETS_DIR / "diagram-database-er.png", "图 1 - MySQL 核心表 ER 图")
    for idx, (table_name, meta) in enumerate(MYSQL_TABLES.items(), start=1):
        doc.add_heading(f"2.{idx} {table_name}", level=2)
        add_paragraph(doc, meta["purpose"])
        add_table(doc, ["字段", "类型", "约束", "说明"], meta["fields"])
        add_table(doc, ["索引名", "字段", "用途"], meta["indexes"])
    doc.add_heading("3. Redis 运行态设计", level=1)
    add_table(doc, ["Key 模式", "数据结构", "用途"], REDIS_ROWS)
    doc.add_heading("4. 设计说明", level=1)
    add_bullet_list(doc, ["reservation.approved_by 当前只保存审批人 ID，未在数据库层建立外键约束。", "meeting_room.equipment_json 使用 JSON 存储设备列表，便于前端按数组直接消费。", "reservation.idx_room_time 和 idx_user_time 已能覆盖房间日历查询与个人近 30 天查询场景。"])
    doc.add_heading("5. 设计建议（显式标注为建议）", level=1)
    add_bullet_list(doc, ["建议：为 reservation.status 增加辅助索引，进一步优化审批列表和状态统计查询。", "建议：如果后续需要审计审批人完整信息，可考虑为 approved_by 增加外键约束或保留快照字段。", "建议：若房间设备维度持续扩展，可将 equipment_json 拆分为独立的 room_equipment 关系表。"])
    output = DOCS_DIR / "05-数据库设计文档.docx"
    return safe_save(doc, output)


def build_api_doc() -> Path:
    doc = create_document("API 接口文档")
    doc.add_heading("1. 接口约定", level=1)
    add_table(doc, ["项目", "约定"], [("服务地址", "开发环境默认 http://127.0.0.1:8080"), ("鉴权方式", "Authorization: Bearer <JWT>"), ("统一响应", "ApiResponse { code, message, data }"), ("时间格式", "LocalDateTime 使用 ISO-8601，如 2026-04-09T09:00:00"), ("在线文档", "Swagger 地址：http://127.0.0.1:8080/doc.html")])
    doc.add_heading("2. 接口清单", level=1)
    add_table(doc, ["分组", "方法", "路径", "权限", "用途"], API_SUMMARY_ROWS)
    doc.add_heading("3. 代表性接口说明", level=1)
    for index, (title, purpose, permission, request_table, request_example, response_example) in enumerate(API_DETAILS, start=1):
        doc.add_heading(f"3.{index} {title}", level=2)
        add_table(doc, ["维度", "说明"], [("用途", purpose), ("权限要求", permission)])
        if request_table:
            add_table(doc, ["参数", "类型", "必填", "说明"], request_table)
        doc.add_paragraph("请求示例")
        add_code_block(doc, request_example)
        doc.add_paragraph("响应示例")
        add_code_block(doc, response_example)
    output = DOCS_DIR / "06-API接口文档.docx"
    return safe_save(doc, output)


def wrap_text_pixels(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.FreeTypeFont, max_width: int) -> list[str]:
    text = str(text).strip()
    if not text:
        return [""]
    lines: list[str] = []
    current = ""
    for char in text:
        candidate = current + char
        width = draw.textbbox((0, 0), candidate, font=font)[2]
        if current and width > max_width:
            lines.append(current.rstrip())
            current = char.lstrip()
        else:
            current = candidate
    if current:
        lines.append(current.rstrip())
    return lines or [text]


def fit_card_text(draw: ImageDraw.ImageDraw, title: str, items: list[str], inner_width: int, body_height: int) -> tuple[ImageFont.FreeTypeFont, list[str], ImageFont.FreeTypeFont, list[str]]:
    title_size = 24
    body_size = 18
    while title_size >= 16:
        title_font = pick_font(title_size, bold=True)
        title_lines = wrap_text_pixels(draw, title, title_font, inner_width)
        if len(title_lines) <= 2:
            break
        title_size -= 1
    else:
        title_font = pick_font(16, bold=True)
        title_lines = wrap_text_pixels(draw, title, title_font, inner_width)

    while body_size >= 12:
        body_font = pick_font(body_size)
        body_lines: list[str] = []
        for item in items:
            body_lines.extend(wrap_text_pixels(draw, item, body_font, inner_width))
        line_height = body_font.size + 6
        if len(body_lines) * line_height <= body_height:
            return title_font, title_lines, body_font, body_lines
        body_size -= 1
    body_font = pick_font(12)
    body_lines = []
    for item in items:
        body_lines.extend(wrap_text_pixels(draw, item, body_font, inner_width))
    return title_font, title_lines, body_font, body_lines


def render_card(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], title: str, items: list[str], *, fill: str, outline: str) -> None:
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=3)
    inner_width = x2 - x1 - 36
    title_font, title_lines, body_font, body_lines = fit_card_text(draw, title, items, inner_width, y2 - y1 - 110)
    draw_multiline_text(draw, (x1 + 18, y1 + 16), title_lines, title_font, fill="#111111", line_gap=4)
    title_height = len(title_lines) * (title_font.size + 4)
    divider_y = y1 + 18 + title_height + 10
    draw.line((x1 + 18, divider_y, x2 - 18, divider_y), fill=outline, width=2)
    draw_multiline_text(draw, (x1 + 18, divider_y + 16), body_lines, body_font, fill="#2f2f2f", line_gap=6)


def draw_poly_arrow(draw: ImageDraw.ImageDraw, points: list[tuple[int, int]], fill: str = "#8b0000", width: int = 4) -> None:
    for start, end in zip(points, points[1:]):
        draw.line((*start, *end), fill=fill, width=width)
    sx, sy = points[-2]
    ex, ey = points[-1]
    angle = math.atan2(ey - sy, ex - sx)
    arrow_len = 16
    spread = math.pi / 8
    left = (ex - arrow_len * math.cos(angle - spread), ey - arrow_len * math.sin(angle - spread))
    right = (ex - arrow_len * math.cos(angle + spread), ey - arrow_len * math.sin(angle + spread))
    draw.polygon([(ex, ey), left, right], fill=fill)


def draw_centered_lines(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], lines: list[str], font: ImageFont.FreeTypeFont, fill: str, line_gap: int = 6) -> None:
    x1, y1, x2, y2 = box
    total_height = len(lines) * font.size + max(0, len(lines) - 1) * line_gap
    y = y1 + max(0, (y2 - y1 - total_height) // 2)
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        width = bbox[2] - bbox[0]
        x = x1 + max(0, (x2 - x1 - width) // 2)
        draw.text((x, y), line, font=font, fill=fill)
        y += font.size + line_gap


def draw_lane_shell(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], title: str, subtitle: str, *, fill: str, outline: str) -> None:
    render_card(draw, xy, title, [subtitle], fill=fill, outline=outline)


def draw_decision_node(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], text: str, *, fill: str, outline: str) -> None:
    x1, y1, x2, y2 = xy
    cx = (x1 + x2) // 2
    cy = (y1 + y2) // 2
    draw.polygon([(cx, y1), (x2, cy), (cx, y2), (x1, cy)], fill=fill, outline=outline)
    font = pick_font(18, bold=True)
    lines = wrap_text_pixels(draw, text, font, int((x2 - x1) * 0.62))
    if len(lines) > 4:
        font = pick_font(16, bold=True)
        lines = wrap_text_pixels(draw, text, font, int((x2 - x1) * 0.66))
    draw_centered_lines(draw, (x1 + 24, y1 + 18, x2 - 24, y2 - 18), lines, font, "#1b1b1b", line_gap=4)


def draw_small_badge(draw: ImageDraw.ImageDraw, center: tuple[int, int], text: str, *, fill: str, outline: str, text_fill: str = "#1b1b1b") -> None:
    cx, cy = center
    font = pick_font(16, bold=True)
    bbox = draw.textbbox((0, 0), text, font=font)
    width = bbox[2] - bbox[0] + 26
    height = 34
    x1 = cx - width // 2
    y1 = cy - height // 2
    x2 = x1 + width
    y2 = y1 + height
    draw.rounded_rectangle((x1, y1, x2, y2), radius=12, fill=fill, outline=outline, width=2)
    draw.text((x1 + 13, y1 + 6), text, font=font, fill=text_fill)


def draw_dashed_lifeline(draw: ImageDraw.ImageDraw, x: int, y1: int, y2: int, *, fill: str = "#99a4b1", width: int = 2, dash: int = 18, gap: int = 10) -> None:
    y = y1
    while y < y2:
        draw.line((x, y, x, min(y + dash, y2)), fill=fill, width=width)
        y += dash + gap


def draw_sequence_participant(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], title: str, *, fill: str = "#f9fbff", outline: str = "#3c78d8") -> None:
    render_card(draw, xy, title, [], fill=fill, outline=outline)


def draw_message(draw: ImageDraw.ImageDraw, y: int, sx: int, ex: int, label: str, *, fill: str = "#fff8eb", outline: str = "#c48a00", arrow_fill: str = "#8b0000") -> None:
    draw_arrow(draw, (sx, y), (ex, y), fill=arrow_fill, width=3)
    mid = (sx + ex) // 2
    span = max(160, min(abs(ex - sx) - 20, 320))
    x1 = max(40, mid - span // 2)
    x2 = x1 + span
    font = pick_font(16)
    lines = wrap_text_pixels(draw, label, font, span - 24)
    height = 18 + len(lines) * (font.size + 2)
    draw.rounded_rectangle((x1, y - 28, x2, y - 28 + height), radius=10, fill=fill, outline=outline, width=2)
    draw_multiline_text(draw, (x1 + 12, y - 22), lines, font, fill="#333333", line_gap=2)


def role_badge_style(tag: str) -> tuple[str, str, str]:
    if tag == "USER":
        return "#e8f2ff", "#3c78d8", "#184a90"
    if tag == "ADMIN":
        return "#fff0e6", "#d97925", "#8b4d0b"
    if tag == "SUPER_ADMIN":
        return "#fdebec", "#c94f5d", "#8b2230"
    return "#f0ecff", "#6b4bb7", "#442c82"


def draw_story_card(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], tag: str, text: str) -> None:
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=16, fill="#ffffff", outline="#cfd6de", width=2)
    badge_fill, badge_outline, badge_text = role_badge_style(tag)
    badge_font = pick_font(14, bold=True)
    badge_bbox = draw.textbbox((0, 0), tag, font=badge_font)
    badge_w = badge_bbox[2] - badge_bbox[0] + 20
    draw.rounded_rectangle((x1 + 12, y1 + 10, x1 + 12 + badge_w, y1 + 36), radius=10, fill=badge_fill, outline=badge_outline, width=2)
    draw.text((x1 + 22, y1 + 15), tag, font=badge_font, fill=badge_text)
    body_font = pick_font(16)
    body_lines = wrap_text_pixels(draw, text, body_font, x2 - x1 - 24)
    draw_multiline_text(draw, (x1 + 12, y1 + 48), body_lines, body_font, fill="#2b2b2b", line_gap=2)


def create_architecture_diagram(path: Path) -> None:
    image, draw = create_canvas(1900, 1180)
    draw.text((70, 42), "系统总体架构图", font=pick_font(42, bold=True), fill="#8b0000")
    draw.text((70, 98), "静态分层视图：统一展示前端、接入安全、业务服务与存储层之间的依赖关系", font=pick_font(22), fill="#444444")

    rows = [(160, 330), (390, 560), (620, 820), (880, 1080)]
    lane_styles = [
        ("浏览器 / 前端层", "页面、路由、状态与请求发起", "#f7efe8", "#b85c38"),
        ("接口与安全接入层", "鉴权过滤链与控制器入口", "#edf6ff", "#3c78d8"),
        ("业务服务层", "业务编排、位图、通知与统计", "#edf7ed", "#3d8b3d"),
        ("存储与运行态层", "MySQL 持久化 + Redis 运行态", "#fff7e9", "#c48a00"),
    ]
    for (y1, y2), (title, subtitle, fill, outline) in zip(rows, lane_styles):
        draw.rounded_rectangle((40, y1, 1860, y2), radius=28, fill="#fcfcfd", outline="#dde3ea", width=2)
        draw_lane_shell(draw, (60, y1 + 18, 250, y2 - 18), title, subtitle, fill=fill, outline=outline)

    card_w = 360
    card_h = 130
    row1_y = 180
    row1_xs = [310, 730, 1150]
    row1_cards = [
        ("浏览器用户", ["PC 浏览器访问系统", "执行预约与审批操作"], "#fff7f2", "#b85c38"),
        ("前端界面层", ["Login / Dashboard / Calendar", "Rooms / Mine / Approvals / Users"], "#f9fbff", "#3c78d8"),
        ("前端基础设施", ["router.beforeEach", "authStore(mrs_auth) + Axios"], "#f9fbff", "#3c78d8"),
    ]
    for x, (title, lines, fill, outline) in zip(row1_xs, row1_cards):
        render_card(draw, (x, row1_y, x + card_w, row1_y + card_h), title, lines, fill=fill, outline=outline)

    row2_y = 410
    render_card(draw, (520, row2_y, 880, row2_y + card_h), "安全过滤链", ["SecurityFilterChain", "JwtAuthFilter 解析 Bearer Token"], fill="#eef5ff", outline="#3c78d8")
    render_card(draw, (980, row2_y, 1340, row2_y + card_h), "控制器入口", ["Auth / Reservation / Admin", "Room / Stats / UserAdmin"], fill="#eef5ff", outline="#3c78d8")

    row3_y = 650
    row3_xs = [310, 730, 1150]
    row3_cards = [
        ("认证服务", ["AuthService + JwtService", "AuthenticationManager 校验后签发 JWT"], "#edf7ed", "#3d8b3d"),
        ("预约服务", ["ReservationService", "时间规则、冲突检测、写库与缓存重建"], "#edf7ed", "#3d8b3d"),
        ("治理与支撑服务", ["ReservationAdmin / MeetingRoom / UserAdmin / Stats", "Bitmap / Lock / Notification / Audit"], "#edf7ed", "#3d8b3d"),
    ]
    for x, (title, lines, fill, outline) in zip(row3_xs, row3_cards):
        render_card(draw, (x, row3_y, x + card_w, row3_y + card_h), title, lines, fill=fill, outline=outline)

    row4_y = 910
    render_card(draw, (620, row4_y, 980, row4_y + card_h), "MySQL 8", ["sys_user / meeting_room / reservation", "权威业务实体与审批结果"], fill="#fff7e9", outline="#c48a00")
    render_card(draw, (1080, row4_y, 1440, row4_y + card_h), "Redis", ["bitmap / roomDayLock / status / maintenances", "notifications / audit / recent reservations"], fill="#fdebec", outline="#c94f5d")

    draw_arrow(draw, (670, 245), (730, 245))
    draw_arrow(draw, (1090, 245), (1150, 245))
    draw_arrow(draw, (1330, 310), (700, 410))
    draw_arrow(draw, (700, 475), (980, 475))
    draw_arrow(draw, (1160, 540), (490, 650))
    draw_arrow(draw, (1160, 540), (910, 650))
    draw_arrow(draw, (1160, 540), (1330, 650))
    draw_arrow(draw, (910, 780), (800, 910))
    draw_arrow(draw, (1330, 780), (1260, 910))
    image.save(path)


def create_auth_flow_diagram(path: Path) -> None:
    image, draw = create_canvas(1700, 1540)
    draw.text((70, 42), "认证与鉴权流程图", font=pick_font(42, bold=True), fill="#8b0000")
    draw.text((70, 98), "控制流视图：登录建立会话、路由守卫、JWT 解析与方法级权限校验", font=pick_font(22), fill="#444444")

    center_x = 760
    box_w = 420
    box_h = 110
    ys = [170, 320, 470, 650, 800, 950, 1110]
    render_card(draw, (center_x - box_w // 2, ys[0], center_x + box_w // 2, ys[0] + box_h), "LoginView 输入凭据", ["/login 输入用户名与密码", "提交认证请求"], fill="#f9fbff", outline="#3c78d8")
    render_card(draw, (center_x - box_w // 2, ys[1], center_x + box_w // 2, ys[1] + box_h), "POST /api/auth/login", ["AuthController 接收请求", "调用 AuthService.login"], fill="#f9fbff", outline="#3c78d8")
    render_card(draw, (center_x - box_w // 2, ys[2], center_x + box_w // 2, ys[2] + box_h), "AuthenticationManager / AuthService", ["校验凭据并检查 enabled 状态"], fill="#f9fbff", outline="#3c78d8")
    draw_decision_node(draw, (635, 590, 885, 760), "凭据正确且账号启用？", fill="#fff7e9", outline="#c48a00")
    render_card(draw, (1000, 615, 1380, 725), "登录失败", ["凭据错误或账号停用", "前端停留在登录页并提示错误"], fill="#fdebec", outline="#c94f5d")
    render_card(draw, (995, ys[3], 1415, ys[3] + box_h), "JwtService + localStorage", ["生成 token / role", "前端写入 mrs_auth"], fill="#edf7ed", outline="#3d8b3d")
    render_card(draw, (center_x - box_w // 2, ys[4], center_x + box_w // 2, ys[4] + 140), "router.beforeEach", ["未登录访问业务页 -> /login", "已登录访问认证页 -> /dashboard", "其余场景继续"], fill="#fff7e9", outline="#c48a00")
    render_card(draw, (1030, ys[4] + 20, 1380, ys[4] + 130), "Axios 请求拦截器", ["从 mrs_auth 读取 token", "写入 Authorization: Bearer"], fill="#f9fbff", outline="#3c78d8")
    render_card(draw, (center_x - box_w // 2, ys[5], center_x + box_w // 2, ys[5] + box_h), "SecurityFilterChain / JwtAuthFilter", ["解析 Bearer Token", "无效 token 按匿名处理"], fill="#f9fbff", outline="#3c78d8")
    draw_decision_node(draw, (635, 1090, 885, 1260), "token 有效？", fill="#fff7e9", outline="#c48a00")
    render_card(draw, (1020, 1115, 1380, 1225), "匿名 / 401", ["无认证上下文", "受保护接口返回 401"], fill="#fdebec", outline="#c94f5d")
    render_card(draw, (center_x - box_w // 2, ys[6] + 190, center_x + box_w // 2, ys[6] + 300), "MethodSecurity", ["按角色注解校验接口权限"], fill="#f9fbff", outline="#3c78d8")
    draw_decision_node(draw, (635, 1340, 885, 1510), "角色权限满足？", fill="#fff7e9", outline="#c48a00")
    render_card(draw, (120, 1365, 480, 1475), "拒绝访问", ["后端返回 403", "前端保留当前上下文"], fill="#fdebec", outline="#c94f5d")
    render_card(draw, (1030, 1365, 1450, 1475), "进入业务页 / 调用业务接口", ["/dashboard、/calendar、/rooms", "以及审批和用户治理接口"], fill="#edf7ed", outline="#3d8b3d")
    render_card(draw, (120, 855, 470, 965), "/login", ["未登录访问业务页时的回退入口"], fill="#fff7e9", outline="#c48a00")
    render_card(draw, (120, 995, 470, 1105), "/dashboard", ["已登录访问 /login 或 /register 时统一跳转"], fill="#fff7e9", outline="#c48a00")

    for start_y, end_y in [(ys[0] + box_h, ys[1]), (ys[1] + box_h, ys[2]), (ys[2] + box_h, 590), (ys[3] + box_h, ys[4]), (ys[4] + 140, ys[5]), (ys[5] + box_h, 1090), (ys[6] + 300, 1340)]:
        draw_arrow(draw, (center_x, start_y), (center_x, end_y))
    draw_poly_arrow(draw, [(885, 675), (995, 675)], fill="#3d8b3d")
    draw.text((910, 640), "是", font=pick_font(16, bold=True), fill="#275d27")
    draw_poly_arrow(draw, [(760, 760), (760, 900)], fill="#3d8b3d")
    draw_poly_arrow(draw, [(995, 725), (1205, 725), (1205, 800)], fill="#3d8b3d")
    draw_poly_arrow(draw, [(635, 675), (470, 675), (470, 670), (470, 800), (1205, 800)], fill="#c94f5d")
    draw.text((560, 640), "否", font=pick_font(16, bold=True), fill="#8b2230")
    draw_poly_arrow(draw, [(970, 930), (1030, 930)], fill="#3c78d8")
    draw_poly_arrow(draw, [(550, 930), (470, 910)], fill="#c48a00")
    draw_poly_arrow(draw, [(550, 970), (470, 1050)], fill="#c48a00")
    draw_poly_arrow(draw, [(1380, 930), (1240, 930), (1240, 950), (970, 950)], fill="#3c78d8")
    draw_poly_arrow(draw, [(885, 1175), (970, 1175), (970, 1160), (1020, 1160)], fill="#c94f5d")
    draw.text((905, 1140), "否", font=pick_font(16, bold=True), fill="#8b2230")
    draw_poly_arrow(draw, [(760, 1260), (760, 1300)], fill="#3d8b3d")
    draw_poly_arrow(draw, [(635, 1425), (480, 1425)], fill="#c94f5d")
    draw.text((585, 1390), "否", font=pick_font(16, bold=True), fill="#8b2230")
    draw_poly_arrow(draw, [(885, 1425), (1030, 1425)], fill="#3d8b3d")
    draw.text((915, 1390), "是", font=pick_font(16, bold=True), fill="#275d27")
    image.save(path)


def create_reservation_sequence_diagram(path: Path) -> None:
    image, draw = create_canvas(2100, 1240)
    draw.text((70, 42), "预约创建时序图", font=pick_font(42, bold=True), fill="#8b0000")
    draw.text((70, 98), "成功路径：收敛为前端、预约服务、房间/位图、锁、MySQL 与通知审计 8 个核心协作对象", font=pick_font(22), fill="#444444")

    participants = [
        ("用户", 120),
        ("CalendarView", 330),
        ("ReservationController", 560),
        ("ReservationService", 800),
        ("MeetingRoomService / Bitmap", 1060),
        ("RedisLockService / Redis", 1320),
        ("ReservationMapper / MySQL", 1600),
        ("Notification / Audit", 1860),
    ]
    for title, x in participants:
        draw_sequence_participant(draw, (x - 95, 170, x + 95, 250), title)
        draw_dashed_lifeline(draw, x, 250, 1140)

    draw_message(draw, 310, 120, 330, "选择房间与时段", fill="#fff8eb", outline="#c48a00")
    draw_message(draw, 380, 330, 560, "POST /api/reservations", fill="#fff8eb", outline="#c48a00")
    draw_message(draw, 450, 560, 800, "create(req)", fill="#fff8eb", outline="#c48a00")
    render_card(draw, (630, 500, 970, 610), "时间规则校验", ["未来时间、单日内", "08:00-18:00、半小时粒度"], fill="#f9fbff", outline="#3c78d8")
    draw_message(draw, 660, 800, 1060, "房间状态 + 维护窗 + 位图预检", fill="#eef5ff", outline="#3c78d8", arrow_fill="#3c78d8")
    draw_message(draw, 730, 1060, 800, "返回可用性结果", fill="#f5fbf4", outline="#3d8b3d", arrow_fill="#3d8b3d")
    draw_message(draw, 820, 800, 1320, "tryLock(roomDayLock)", fill="#fff8eb", outline="#c48a00")
    draw_message(draw, 890, 1320, 800, "lock acquired", fill="#f5fbf4", outline="#3d8b3d", arrow_fill="#3d8b3d")
    draw_message(draw, 980, 800, 1600, "ensureNoConflict() + insert reservation", fill="#fff8eb", outline="#c48a00")
    render_card(draw, (1430, 1020, 1860, 1130), "状态说明", ["A-101 或 requireApproval=true", "初始状态为 PENDING，否则为 APPROVED"], fill="#fff7e9", outline="#c48a00")
    draw_message(draw, 1050, 1600, 800, "返回 ReservationResp", fill="#f5fbf4", outline="#3d8b3d", arrow_fill="#3d8b3d")
    draw_message(draw, 1120, 800, 1060, "rebuildDay + evictUserCache", fill="#eef5ff", outline="#3c78d8", arrow_fill="#3c78d8")
    draw_message(draw, 1190, 800, 1860, "pushToUser() + logForUser()", fill="#eef5ff", outline="#3c78d8", arrow_fill="#3c78d8")
    draw_message(draw, 1260, 800, 1320, "unlock(roomDayLock)", fill="#fff8eb", outline="#c48a00")
    draw_message(draw, 1330, 800, 560, "ReservationResp", fill="#f5fbf4", outline="#3d8b3d", arrow_fill="#3d8b3d")
    draw_message(draw, 1400, 560, 330, "ApiResponse<ReservationResp>", fill="#f5fbf4", outline="#3d8b3d", arrow_fill="#3d8b3d")
    draw_message(draw, 1470, 330, 120, "刷新日历并提示结果", fill="#f5fbf4", outline="#3d8b3d", arrow_fill="#3d8b3d")
    image.save(path)


def create_story_map_diagram(path: Path) -> None:
    image, draw = create_canvas(1460, 2100)
    draw.text((60, 42), "用户故事地图（Backbone / Activities / Stories / Slices）", font=pick_font(38, bold=True), fill="#8b0000")
    draw.text((60, 94), "为适配竖版文档，采用 4+4 分段展示主干活动；每段仍保持 backbone、tasks 与 3 个优先级切片", font=pick_font(22), fill="#444444")
    draw.text((1010, 110), "角色标签：USER / ADMIN / SUPER_ADMIN", font=pick_font(18, bold=True), fill="#444444")

    columns = [
        ("登录进入", "输入凭据\n建立会话"),
        ("查看看板", "进入 dashboard\n查看提醒与风险"),
        ("查找会议室", "筛选资源\n查看日历占用"),
        ("发起预约", "选房间时段\n提交预约"),
        ("管理个人预约", "查看近期记录\n取消或删除"),
        ("审批协同", "处理待审批\n追踪结果"),
        ("会议室治理", "维护房间\n状态与维护窗"),
        ("用户治理", "查看用户\n角色与账号状态"),
    ]
    slices = [
        (
            "P0 实验核心",
            {
                "登录进入": [("USER/ADMIN/SUPER_ADMIN", "登录系统并统一进入 /dashboard")],
                "查看看板": [("USER", "查看欢迎语与个人提醒"), ("ADMIN", "查看待审批数量与维护风险")],
                "查找会议室": [("USER", "按容量/设备筛选并查看日历占用")],
                "发起预约": [("USER", "提交单次预约请求")],
                "管理个人预约": [("USER", "查看我的预约并取消预约")],
                "审批协同": [("ADMIN", "查看待审批列表并批准/驳回")],
            },
        ),
        (
            "P1 当前已实现增强",
            {
                "登录进入": [("USER", "注册新账号并进入系统")],
                "查看看板": [("SUPER_ADMIN", "查看全局资源态势与周趋势"), ("ADMIN", "从通知中心查看审批/维护消息")],
                "查找会议室": [("ADMIN/SUPER_ADMIN", "结合状态与维护窗判断资源可用性")],
                "发起预约": [("USER", "按周批量创建周期预约"), ("USER", "获取冲突建议与可用时段提示")],
                "管理个人预约": [("USER/ADMIN/SUPER_ADMIN", "删除预约并查看近 30 天个人记录")],
                "审批协同": [("ADMIN", "查看已审批记录与审计备注")],
            },
        ),
        (
            "P2 治理扩展",
            {
                "查看看板": [("SUPER_ADMIN", "识别高风险房间与维护压力")],
                "审批协同": [("SUPER_ADMIN", "将已审结果撤回到 PENDING")],
                "会议室治理": [("SUPER_ADMIN", "新增/编辑/删除会议室"), ("SUPER_ADMIN", "维护房间状态与维护时段")],
                "用户治理": [("ADMIN/SUPER_ADMIN", "查看用户列表并修改角色"), ("ADMIN/SUPER_ADMIN", "启用或停用账号")],
            },
        ),
    ]

    panel_specs = [(0, 4, 170), (4, 8, 1080)]
    start_x = 220
    col_w = 250
    col_gap = 18
    header_h = 82
    activity_h = 96
    slice_h = 210
    card_h = 78

    for start, end, panel_y in panel_specs:
        panel_title = f"主干活动段 {start + 1}-{end}"
        render_card(draw, (60, panel_y, 180, panel_y + header_h), panel_title, [], fill="#fff7f2", outline="#b85c38")
        for offset, (title, activity) in enumerate(columns[start:end]):
            x1 = start_x + offset * (col_w + col_gap)
            x2 = x1 + col_w
            render_card(draw, (x1, panel_y, x2, panel_y + header_h), title, [], fill="#f4f8ff", outline="#3c78d8")
            render_card(draw, (x1, panel_y + 98, x2, panel_y + 98 + activity_h), "关键任务", activity.split("\n"), fill="#ffffff", outline="#cfd6de")

        for slice_index, (slice_title, stories) in enumerate(slices):
            y1 = panel_y + 220 + slice_index * slice_h
            y2 = y1 + slice_h - 16
            draw_lane_shell(draw, (60, y1, 180, y2), slice_title, "同一横切片代表同一优先级", fill="#fff2eb", outline="#b85c38")
            for offset, (title, _) in enumerate(columns[start:end]):
                x1 = start_x + offset * (col_w + col_gap)
                x2 = x1 + col_w
                draw.rounded_rectangle((x1, y1, x2, y2), radius=18, fill="#fafbfc", outline="#d9e0e6", width=2)
                cards = stories.get(title, [])
                cursor = y1 + 12
                for tag, text in cards[:2]:
                    draw_story_card(draw, (x1 + 12, cursor, x2 - 12, cursor + card_h), tag, text)
                    cursor += card_h + 10
    image.save(path)


def create_database_er_diagram(path: Path) -> None:
    image, draw = create_canvas(1840, 1000)
    draw.text((70, 42), "MySQL 核心表 ER 图", font=pick_font(42, bold=True), fill="#8b0000")
    draw.text((70, 98), "实体视图：统一尺寸的实体卡片，仅表达字段、主外键与 1:N 关系", font=pick_font(22), fill="#444444")

    user_box = (80, 220, 520, 760)
    reservation_box = (700, 220, 1140, 760)
    room_box = (1320, 220, 1760, 760)
    render_card(draw, user_box, "sys_user", ["PK id : bigint", "username : varchar(64) UNIQUE", "password_hash : varchar(255)", "role : varchar(32)", "enabled : tinyint", "created_at / updated_at : datetime"], fill="#f9fbff", outline="#3c78d8")
    render_card(draw, reservation_box, "reservation", ["PK id : bigint", "FK user_id -> sys_user.id", "FK room_id -> meeting_room.id", "start_time / end_time : datetime", "status : varchar(32)", "reason / admin_comment : varchar(255)", "approved_by / approved_at", "created_at / updated_at : datetime"], fill="#fff7e9", outline="#c48a00")
    render_card(draw, room_box, "meeting_room", ["PK id : bigint", "name : varchar(128) UNIQUE", "capacity : int", "equipment_json : json", "require_approval : tinyint", "created_at / updated_at : datetime"], fill="#edf7ed", outline="#3d8b3d")

    draw.line((520, 490, 700, 490), fill="#8b0000", width=4)
    draw.line((1140, 490, 1320, 490), fill="#8b0000", width=4)
    draw_small_badge(draw, (585, 455), "1", fill="#f9fbff", outline="#3c78d8")
    draw_small_badge(draw, (650, 455), "N", fill="#fff7e9", outline="#c48a00")
    draw_small_badge(draw, (1190, 455), "N", fill="#fff7e9", outline="#c48a00")
    draw_small_badge(draw, (1255, 455), "1", fill="#edf7ed", outline="#3d8b3d")
    render_card(draw, (590, 820, 1250, 930), "关系说明", ["sys_user 1:N reservation", "meeting_room 1:N reservation", "Redis 不属于 ER 图范围"], fill="#fffdf7", outline="#c48a00")
    image.save(path)


def set_cell_text(cell, text: str, *, bold: bool = False, size: float = 12, color: RGBColor | None = None, align: int = WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if color:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def apply_run_font(run, ascii_font: str, east_asia_font: str, size: float, *, bold: bool | None = None) -> None:
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:ascii"), ascii_font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), ascii_font)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)
    run._element.rPr.rFonts.set(qn("w:cs"), ascii_font)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_page(doc: Document, *, landscape: bool = False, section_index: int = 0) -> None:
    section = doc.sections[section_index]
    section.page_width = Cm(29.7 if landscape else 21)
    section.page_height = Cm(21 if landscape else 29.7)
    section.orientation = WD_ORIENTATION.LANDSCAPE if landscape else WD_ORIENTATION.PORTRAIT
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)


def set_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "SimSun"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "SimSun")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "SimSun")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal._element.rPr.rFonts.set(qn("w:cs"), "SimSun")
    normal.font.size = Pt(12)
    normal.font.bold = False
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Cm(0.84)
    style_specs = [("Heading 1", 16, True), ("Heading 2", 14, True), ("Heading 3", 12, True)]
    for style_name, size, bold in style_specs:
        style = ensure_paragraph_style(doc, style_name)
        style.font.name = "SimSun"
        style._element.rPr.rFonts.set(qn("w:ascii"), "SimSun")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "SimSun")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style._element.rPr.rFonts.set(qn("w:cs"), "SimSun")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def ensure_paragraph_style(doc: Document, style_name: str):
    try:
        return doc.styles[style_name]
    except KeyError:
        style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Normal"]
        return style


def create_document(title: str, landscape: bool = False) -> Document:
    doc = Document(TEMPLATE_PATH) if TEMPLATE_PATH.exists() else Document()
    set_page(doc, landscape=False, section_index=0)
    set_styles(doc)
    doc.core_properties.title = title
    doc.core_properties.subject = PROJECT_NAME
    doc.core_properties.author = "Codex"
    doc.core_properties.comments = "课程设计文档"
    add_cover(doc, title)
    new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    set_page(doc, landscape=landscape, section_index=len(doc.sections) - 1)
    return doc


def resolve_cover_title(title: str) -> str:
    return COVER_TITLE_MAP.get(title, title)


def add_cover(doc: Document, title: str) -> None:
    if TEMPLATE_PATH.exists():
        while len(doc.paragraphs) > 16:
            delete_paragraph(doc.paragraphs[-1])
        title_para = doc.paragraphs[4]
        while title_para.runs:
            title_para.runs[-1]._element.getparent().remove(title_para.runs[-1]._element)
        run = title_para.add_run("题    目：" + resolve_cover_title(title))
        apply_run_font(run, "Microsoft YaHei", "微软雅黑", 22, bold=True)

        date_para = doc.paragraphs[11]
        while date_para.runs:
            date_para.runs[-1]._element.getparent().remove(date_para.runs[-1]._element)
        run = date_para.add_run("完成日期：" + TODAY)
        apply_run_font(run, "Microsoft YaHei", "微软雅黑", 22, bold=True)
        return

    p = doc.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("   课程设计")
    apply_run_font(run, "Microsoft YaHei", "微软雅黑", 26, bold=True)
    rows = [
        f"课程名称：私有云平台架构与实践",
        f"题    目：{resolve_cover_title(title)}",
        "学    院：示范性软件学院",
        "学生姓名：张义南",
        "学    号：2300770217",
        "年    级：23级",
        "专业班级：软工2303",
        "任课教师：丁玺润",
        f"完成日期：{TODAY}",
    ]
    for text in rows:
        para = doc.add_paragraph()
        para.paragraph_format.line_spacing = 1.5
        run = para.add_run(text)
        apply_run_font(run, "Microsoft YaHei", "微软雅黑", 22, bold=True)


def add_paragraph(doc: Document, text: str) -> None:
    para = doc.add_paragraph(text)
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.first_line_indent = Cm(0.84)


def add_bullet_list(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        try:
            para = doc.add_paragraph(style="List Bullet")
        except KeyError:
            para = doc.add_paragraph()
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.first_line_indent = Cm(0)
        run = para.add_run(item)
        run.font.name = "SimSun"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(12)


def add_image(doc: Document, path: Path, caption: str, width_cm: float = 14.2) -> None:
    if not path.exists():
        add_paragraph(doc, f"图资源缺失：{path.name}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True
    section = doc.sections[-1]
    max_width_cm = (section.page_width - section.left_margin - section.right_margin) / Cm(1)
    p.add_run().add_picture(str(path), width=Cm(min(width_cm, max_width_cm - 0.4)))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Cm(0)
    cap.paragraph_format.line_spacing = 1.5
    cap.paragraph_format.keep_with_next = True
    cap.paragraph_format.keep_together = True
    run = cap.add_run(caption)
    apply_run_font(run, "SimHei", "黑体", 10.5, bold=False)


def add_code_block(doc: Document, data: dict | list | None) -> None:
    if data is None:
        add_paragraph(doc, "无")
        return
    for line in json.dumps(data, ensure_ascii=False, indent=2).splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.35
        run = p.add_run(line)
        apply_run_font(run, "SimSun", "宋体", 12)


def normalize_document(doc: Document) -> None:
    heading_fonts = {
        "Heading 1": (16, True),
        "Heading 2": (14, True),
        "Heading 3": (12, True),
    }
    for para in doc.paragraphs:
        style_name = para.style.name if para.style is not None else ""
        if style_name in heading_fonts:
            size, bold = heading_fonts[style_name]
            for run in para.runs:
                apply_run_font(run, "SimSun", "宋体", size, bold=bold)


def safe_save(doc: Document, output: Path) -> Path:
    normalize_document(doc)
    try:
        doc.save(output)
        return output
    except PermissionError:
        for index in range(1, 100):
            suffix = "-课程设计版" if index == 1 else f"-课程设计版-v{index}"
            fallback = output.with_name(f"{output.stem}{suffix}{output.suffix}")
            try:
                doc.save(fallback)
                return fallback
            except PermissionError:
                continue
        raise


def main() -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    create_diagram_assets()
    outputs = [build_srs(), build_prototype_doc(), build_story_map_doc(), build_architecture_doc(), build_database_doc(), build_api_doc()]
    print("已生成文档：")
    for output in outputs:
        print(output)


if __name__ == "__main__":
    main()
